from __future__ import annotations

import csv
import io
import re
from pathlib import Path

SUPPORTED_EXTENSIONS = (
    ".txt", ".md", ".pdf", ".docx", ".xlsx", ".csv", ".html", ".htm",
)


def chunk_text(
    text: str,
    chunk_size: int = 800,
    overlap: int = 128,
) -> list[str]:
    paragraphs = _split_paragraphs(text)
    chunks = _merge_paragraphs(paragraphs, chunk_size, overlap)
    return [c for c in chunks if c.strip()]


def _split_paragraphs(text: str) -> list[str]:
    blocks = re.split(r"\n\n+", text)
    return [b.strip() for b in blocks if b.strip()]


def _merge_paragraphs(
    paragraphs: list[str],
    chunk_size: int,
    overlap: int,
) -> list[str]:
    chunks: list[str] = []
    buffer: list[str] = []
    current_size = 0

    for para in paragraphs:
        para_chars = len(para)

        if current_size + para_chars > chunk_size and buffer:
            chunks.append("\n\n".join(buffer))
            overlap_chars = 0
            overlap_buf: list[str] = []
            for p in reversed(buffer):
                pc = len(p)
                if overlap_chars + pc <= overlap:
                    overlap_buf.insert(0, p)
                    overlap_chars += pc
                else:
                    break
            buffer = overlap_buf
            current_size = overlap_chars

        buffer.append(para)
        current_size += para_chars

    if buffer:
        chunks.append("\n\n".join(buffer))

    return chunks


def chunk_file(
    filepath: str | Path,
    chunk_size: int = 512,
    overlap: int = 64,
) -> list[str]:
    filepath = Path(filepath)
    suffix = filepath.suffix.lower()

    if suffix in (".txt", ""):
        text = filepath.read_text(encoding="utf-8")
        return chunk_text(text, chunk_size, overlap)

    if suffix == ".md":
        text = _read_markdown(filepath)
        return chunk_text(text, chunk_size, overlap)

    if suffix == ".pdf":
        try:
            import pymupdf  # noqa: F401
        except ImportError:
            raise ImportError(
                "pymupdf not installed. Run: pip install pymupdf"
            ) from None
        return _read_pdf(filepath, chunk_size, overlap)

    if suffix == ".docx":
        try:
            import docx  # noqa: F401
        except ImportError:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            ) from None
        return _read_docx(filepath, chunk_size, overlap)

    if suffix == ".csv":
        return _read_csv(filepath, chunk_size, overlap)

    if suffix == ".xlsx":
        try:
            import openpyxl  # noqa: F401
        except ImportError:
            raise ImportError(
                "openpyxl not installed. Run: pip install openpyxl"
            ) from None
        return _read_xlsx(filepath, chunk_size, overlap)

    if suffix in (".html", ".htm"):
        return _read_html(filepath, chunk_size, overlap)

    raise ValueError(f"unsupported file type: {suffix}")


def _read_markdown(filepath: Path) -> str:
    text = filepath.read_text(encoding="utf-8")
    text = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"`{1,3}[^`]*`{1,3}", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"[*_]{1,3}", "", text)
    return text


def _read_pdf(filepath: Path, chunk_size: int, overlap: int) -> list[str]:
    import pymupdf

    doc = pymupdf.open(str(filepath))
    text_parts: list[str] = []
    for page in doc:
        page_text = page.get_text()
        if page_text.strip():
            text_parts.append(page_text.strip())
    doc.close()
    full = "\n\n".join(text_parts)
    return chunk_text(full, chunk_size, overlap)


def _read_docx(filepath: Path, chunk_size: int, overlap: int) -> list[str]:
    import docx

    document = docx.Document(str(filepath))
    text_parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text.strip():
                text_parts.append(row_text)
    full = "\n\n".join(text_parts)
    return chunk_text(full, chunk_size, overlap)


def _read_csv(filepath: Path, chunk_size: int, overlap: int) -> list[str]:
    text = filepath.read_text(encoding="utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    parts = [" | ".join(row) for row in rows if any(cell.strip() for cell in row)]
    return chunk_text("\n".join(parts), chunk_size, overlap)


def _read_xlsx(filepath: Path, chunk_size: int, overlap: int) -> list[str]:
    import openpyxl

    wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    wb.close()
    return chunk_text("\n".join(parts), chunk_size, overlap)


def _read_html(filepath: Path, chunk_size: int, overlap: int) -> list[str]:
    text = filepath.read_text(encoding="utf-8", errors="replace")
    text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"&lt;", "<", text)
    text = re.sub(r"&gt;", ">", text)
    text = re.sub(r"&quot;", '"', text)
    text = re.sub(r"&#?\w+;", " ", text)
    text = re.sub(r"\s+", " ", text)
    return chunk_text(text, chunk_size, overlap)
